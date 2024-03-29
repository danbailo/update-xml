from docx import Document
from docx.shared import Inches
import xmltodict
import re
import os

class Xml:
    def __init__(self, file = None, cnpj = None):        
        self.__xml = file
        self.__cnpj = cnpj
        self.__report = Document()
        self.__report.add_heading("XMLs que não sofreram alterações",level=0)

    def get_xml(self):
        if os.path.isfile(self.__xml):
            if self.__xml[-4:].lower() == ".xml":
                self.__isfile = True
                self.__isdir = False
                return [self.__xml.split("/")[-1]]
            print("Por favor, entre com um arquivo .xml!")
            exit(-1)
        if os.path.isdir(self.__xml):
            xmls = [xml for xml in os.listdir(self.__xml) if xml[-4:].lower()==".xml"]
            if len(xmls) == 0: return None
            self.__isfile = False
            self.__isdir = True
            return xmls

    def __check_cnpj(self, cnpj, doc):
        if cnpj == doc['nfeProc']['NFe']['infNFe']["emit"]["CNPJ"]:
            return True
        return False

    def get_cnpj(self):
        with open(self.__cnpj,"r") as csv:
            all_cnpj = [re.sub(pattern=r"\D", repl="", string=cnpj) for cnpj in csv.readlines()]
        return set(all_cnpj)
  
    def update_fields(self):        
        for xml in self.get_xml():
            if len(self.get_xml()) == 1:
                fd = open(os.path.join(self.__xml))
            else : 
                fd = open(os.path.join(self.__xml,xml))

            try:
                text = fd.read()
            except Exception:
                fd.close()
                del fd
                if len(self.get_xml()) == 1:
                    fd = open(os.path.join(self.__xml), encoding="utf-8")
                else: 
                    fd = open(os.path.join(self.__xml,xml), encoding="utf-8")
                text = fd.read()

            doc = xmltodict.parse(text)
            fd.close()
            for cnpj in self.get_cnpj():
                if self.__check_cnpj(cnpj, doc):
                    state = 1
                    break
                else:
                    state = 0
            if state == 0: continue        
            print('\nO CNPJ "{cnpj}" foi encontrado no arquivo "{xml}"'.format(cnpj=cnpj, xml=xml))

            if isinstance(doc['nfeProc']['NFe']["infNFe"]["det"], list):
                for i in range(len(doc['nfeProc']['NFe']["infNFe"]["det"])):
                    new_doc = doc.copy()
                    try:
                        vICMS = new_doc['nfeProc']['NFe']['infNFe']['det'][i]['imposto']["ICMS"]["ICMS00"]["vICMS"]
                        #REGRA 1 E REGRA 3
                        vBC_PIS = new_doc['nfeProc']['NFe']['infNFe']['det'][i]['imposto']["PIS"]["PISAliq"]["vBC"]
                        new_vBC_PIS = str(round(eval(vBC_PIS+'-'+vICMS),2))
                        new_doc['nfeProc']['NFe']['infNFe']['det'][i]['imposto']["PIS"]["PISAliq"]["vBC"] = new_vBC_PIS

                        vBC_COFINS = new_doc['nfeProc']['NFe']['infNFe']['det'][i]['imposto']["COFINS"]["COFINSAliq"]["vBC"]
                        new_vBC_COFINS = str(round(eval(vBC_COFINS+'-'+vICMS),2))
                        new_doc['nfeProc']['NFe']['infNFe']['det'][i]['imposto']["COFINS"]["COFINSAliq"]["vBC"] = new_vBC_COFINS
                        
                        #REGRA 2
                        pPIS = new_doc['nfeProc']['NFe']['infNFe']['det'][i]['imposto']["PIS"]["PISAliq"]["pPIS"]
                        new_vPIS = str(round(eval(pPIS+'*'+new_vBC_PIS),2))
                        new_vPIS = str(round(eval(new_vPIS+'*0.01'),2))
                        new_doc['nfeProc']['NFe']['infNFe']['det'][i]['imposto']["PIS"]["PISAliq"]["vPIS"] = new_vPIS
                        
                        #REGRA 4
                        pCOFINS = new_doc['nfeProc']['NFe']['infNFe']['det'][i]['imposto']["COFINS"]["COFINSAliq"]["pCOFINS"]
                        new_vCOFINS = str(round(eval(pCOFINS+'*'+new_vBC_COFINS),2))
                        new_vCOFINS = str(round(eval(new_vCOFINS+'*0.01'),2))
                        new_doc['nfeProc']['NFe']['infNFe']['det'][i]['imposto']["COFINS"]["COFINSAliq"]["vCOFINS"] = new_vCOFINS
                        
                        #REGRA 5
                        new_doc['nfeProc']['NFe']['infNFe']["total"]["ICMSTot"]["vPIS"] = new_vPIS
                        new_doc['nfeProc']['NFe']['infNFe']["total"]["ICMSTot"]["vCOFINS"] = new_vCOFINS
                    except Exception:
                        natop = doc["nfeProc"]["NFe"]["infNFe"]["ide"]["natOp"]
                        vicmstot = doc["nfeProc"]["NFe"]["infNFe"]["total"]["ICMSTot"]["vICMS"]
                        vbctot = doc["nfeProc"]["NFe"]["infNFe"]["total"]["ICMSTot"]["vBC"]
                        cst = ''
                        csosn = ''
                        print("Campo natureza da operação:{natop}".format(natop=natop))
                        print("ICMSTot: vICMS = {vicmstot}".format(vicmstot=vicmstot))
                        print("ICMSTot: vBC = {vbctot}".format(vbctot=vbctot))                        
                        
                        err_doc = doc['nfeProc']['NFe']["infNFe"]["det"]
                        if isinstance(err_doc, list):
                            for i in range(len(err_doc)):
                                for k in err_doc[i]['imposto']["ICMS"].keys():
                                    if "CST" in err_doc[i]['imposto']["ICMS"][k].keys():
                                        cst = err_doc[i]['imposto']["ICMS"][k]["CST"]
                                        print("CST: {cst}".format(cst=cst))
                                    if "CSOSN" in err_doc[i]['imposto']["ICMS"][k].keys():
                                        csosn = err_doc[i]['imposto']["ICMS"][k]["CSOSN"]
                                        print("CSOSN: {csosn}".format(csosn=csosn))
                        self.__report.add_heading("Arquivo: {xml}".format(xml=xml),level=1)
                        self.__report.add_paragraph("Campo natureza da operação:{natop}".format(natop=natop))
                        self.__report.add_paragraph("ICMSTot: vICMS = {vicmstot}".format(vicmstot=vicmstot))
                        self.__report.add_paragraph("ICMSTot: vBC = {vbctot}".format(vbctot=vbctot))
                        if cst:
                            self.__report.add_paragraph("CST: {cst}".format(cst=cst))
                        if csosn:
                            self.__report.add_paragraph("CSOSN: {csosn}".format(csosn=csosn))
                        state = -1
                        break
            else:
                new_doc = doc.copy()
                try:
                    vICMS = new_doc['nfeProc']['NFe']['infNFe']['det']['imposto']["ICMS"]["ICMS00"]["vICMS"]
                    
                    #REGRA 1 E REGRA 3
                    vBC_PIS = new_doc['nfeProc']['NFe']['infNFe']['det']['imposto']["PIS"]["PISAliq"]["vBC"]
                    new_vBC_PIS = str(round(eval(vBC_PIS+'-'+vICMS),2))
                    new_doc['nfeProc']['NFe']['infNFe']['det']['imposto']["PIS"]["PISAliq"]["vBC"] = new_vBC_PIS

                    vBC_COFINS = new_doc['nfeProc']['NFe']['infNFe']['det']['imposto']["COFINS"]["COFINSAliq"]["vBC"]
                    new_vBC_COFINS = str(round(eval(vBC_COFINS+'-'+vICMS),2))
                    new_doc['nfeProc']['NFe']['infNFe']['det']['imposto']["COFINS"]["COFINSAliq"]["vBC"] = new_vBC_COFINS

                    #REGRA 2
                    pPIS = new_doc['nfeProc']['NFe']['infNFe']['det']['imposto']["PIS"]["PISAliq"]["pPIS"]
                    new_vPIS = str(round(eval(pPIS+'*'+new_vBC_PIS),2))
                    new_vPIS = str(round(eval(new_vPIS+'*0.01'),2))
                    new_doc['nfeProc']['NFe']['infNFe']['det']['imposto']["PIS"]["PISAliq"]["vPIS"] = new_vPIS

                    #REGRA 4
                    pCOFINS = new_doc['nfeProc']['NFe']['infNFe']['det']['imposto']["COFINS"]["COFINSAliq"]["pCOFINS"]
                    new_vCOFINS = str(round(eval(pCOFINS+'*'+new_vBC_COFINS),2))
                    new_vCOFINS = str(round(eval(new_vCOFINS+'*0.01'),2))
                    new_doc['nfeProc']['NFe']['infNFe']['det']['imposto']["COFINS"]["COFINSAliq"]["vCOFINS"] = new_vCOFINS

                    #REGRA 5
                    new_doc['nfeProc']['NFe']['infNFe']["total"]["ICMSTot"]["vPIS"] = new_vPIS
                    new_doc['nfeProc']['NFe']['infNFe']["total"]["ICMSTot"]["vCOFINS"] = new_vCOFINS
                except Exception:
                    natop = doc["nfeProc"]["NFe"]["infNFe"]["ide"]["natOp"]
                    vicmstot = doc["nfeProc"]["NFe"]["infNFe"]["total"]["ICMSTot"]["vICMS"]
                    vbctot = doc["nfeProc"]["NFe"]["infNFe"]["total"]["ICMSTot"]["vBC"]
                    cst = ''
                    csosn = ''
                    print("Campo natureza da operação: {natop}".format(natop=natop))
                    print("ICMSTot: vICMS = {vicmstot}".format(vicmstot=vicmstot))
                    print("ICMSTot: vBC = {vbctot}".format(vbctot=vbctot))       

                    err_doc = doc['nfeProc']['NFe']["infNFe"]["det"]
                    err_ICMS = err_doc['imposto']["ICMS"]
                    for k in err_ICMS.keys():
                        if "CST" in err_ICMS[k].keys():
                            cst = err_ICMS[k]["CST"]
                            print("CST: {cst}".format(cst=cst))
                        if "CSOSN" in err_ICMS[k].keys():
                            csosn = err_ICMS[k]["CSOSN"]
                            print("CSOSN: {csosn}".format(csosn=csosn))
                    
                    self.__report.add_heading("Arquivo: {xml}".format(xml=xml),level=1)
                    self.__report.add_paragraph("Campo natureza da operação: {natop}".format(natop=natop))
                    self.__report.add_paragraph("ICMSTot: vICMS = {vicmstot}".format(vicmstot=vicmstot))
                    self.__report.add_paragraph("ICMSTot: vBC = {vbctot}".format(vbctot=vbctot))
                    if cst:
                        self.__report.add_paragraph("CST: {cst}".format(cst=cst))
                    if csosn:
                        self.__report.add_paragraph("CSOSN: {csosn}".format(csosn=csosn))
                    state = -1
            if state != -1:
                if not os.path.exists(os.path.join("..","output")):
                    os.mkdir(os.path.join("..","output"))
                if self.__isfile:
                    xml_splitted = xml[:-4].split("\\")
                    xml = xml_splitted[-1]
                    result_file = open(os.path.join(".","..","output",xml+" - ALTERADO.xml"), 'w')
                    xml = xml+".xml"
                else:
                    result_file = open(os.path.join(".","..","output",xml[:-4]+" - ALTERADO.xml"), 'w')
                print(f'O arquivo "{xml}" foi alterado com sucesso!\n')
                result_file.write(xmltodict.unparse(new_doc,full_document=False))
                result_file.close()
        if state==0: print("CNPJ emitente não encontrado neste XML!".format(cnpj=cnpj))
        try:
            self.__report.save(os.path.join("..","Relatório.docx"))
            print("O relatório dos XMLs que não foram alterados foi gerado com sucesso!")
        except Exception:
            print("Por favor, feche o documento Word e execute novamente o programa para gerar um novo relatório!")
            exit(-1)